import io
import re
import fitz
import pytesseract
from docx import Document as DocxDocument
from PIL import Image
from pillow_heif import register_heif_opener

register_heif_opener()

SUPPORTED_IMAGE_FORMATS = {"JPEG", "PNG", "TIFF", "BMP", "WEBP", "HEIF"}


def _looks_like_text(file_bytes: bytes) -> bool:
    if not file_bytes:
        return False
    try:
        sample = file_bytes[:8192].decode("utf-8")
    except UnicodeDecodeError:
        return False
    return "\x00" not in sample


def detect_file_type(file_bytes: bytes) -> str | None:
    """Sniff the actual file type from content, not the filename.

    Returns "pdf", "image", "docx", "text", or None if unsupported.
    """
    if file_bytes[:5] == b"%PDF-":
        return "pdf"
    if file_bytes[:4] == b"PK\x03\x04":
        try:
            DocxDocument(io.BytesIO(file_bytes))
            return "docx"
        except Exception:
            pass
    try:
        img = Image.open(io.BytesIO(file_bytes))
        img.verify()
        if img.format in SUPPORTED_IMAGE_FORMATS:
            return "image"
    except Exception:
        pass
    if _looks_like_text(file_bytes):
        return "text"
    return None


async def extract_text_from_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    texts = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            texts.append(text)
        else:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            texts.append(pytesseract.image_to_string(img, lang="eng"))
    return "\n".join(texts)


async def extract_text_from_image(file_bytes: bytes) -> str:
    img = Image.open(io.BytesIO(file_bytes))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    return pytesseract.image_to_string(img, lang="eng")


async def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


async def extract_text_from_plain(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8")


def extract_lab_markers(text: str) -> list[dict]:
    markers = []
    pattern = re.compile(
        r"([A-Za-z][A-Za-z0-9\s\-\.]{1,40}?)\s*[:\|]?\s*"
        r"([<>]?\d+\.?\d*)\s*"
        r"([a-zA-Z/%μ]+)?\s*"
        r"(?:Ref(?:erence)?\s*Range?:?\s*([\d\.\-\s]+[a-zA-Z/%μ]*))?",
        re.MULTILINE,
    )
    for m in pattern.finditer(text):
        name = m.group(1).strip()
        if len(name) < 2 or len(name) > 50:
            continue
        markers.append({
            "name": name,
            "value": m.group(2),
            "unit": m.group(3) or "",
            "reference_range": m.group(4) or "",
        })
    return markers[:50]
